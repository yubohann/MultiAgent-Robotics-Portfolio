from __future__ import annotations

import os
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


def target_report() -> Path:
    env_path = os.environ.get("REPORT_TARGET")
    if env_path:
        return Path(env_path)
    repo = Path.home() / "supermarket-management-system" / "supermarket-management-system"
    return next((repo / "reports" / "system-analysis-design").glob("*20260515.docx"))


def insert_after(paragraph: Paragraph) -> Paragraph:
    element = OxmlElement("w:p")
    paragraph._p.addnext(element)
    return Paragraph(element, paragraph._parent)


def set_font(run, size=11, bold=False, font_name="宋体") -> None:
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def add_paragraph_after(cursor: Paragraph, text: str, bold=False, first_line=True) -> Paragraph:
    para = insert_after(cursor)
    fmt = para.paragraph_format
    fmt.line_spacing = 1.5
    fmt.space_after = Pt(6)
    if first_line:
        fmt.first_line_indent = Pt(22)
    run = para.add_run(text)
    set_font(run, size=11, bold=bold, font_name="黑体" if bold else "宋体")
    return para


def add_heading_after(cursor: Paragraph, text: str) -> Paragraph:
    para = insert_after(cursor)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt = para.paragraph_format
    fmt.space_before = Pt(8)
    fmt.space_after = Pt(4)
    run = para.add_run(text)
    set_font(run, size=12, bold=True, font_name="黑体")
    return para


def find_heading(doc: Document, exact_text: str) -> Paragraph:
    for para in doc.paragraphs:
        if para.text.strip() == exact_text:
            return para
    raise ValueError(f"Heading not found: {exact_text}")


def insert_before_heading(doc: Document, heading_text: str, blocks: list[tuple[str, list[str]]]) -> None:
    heading = find_heading(doc, heading_text)
    cursor = heading
    inserted: list[Paragraph] = []
    for title, paragraphs in blocks:
        h = add_heading_after(cursor, title)
        inserted.append(h)
        cursor = h
        for text in paragraphs:
            p = add_paragraph_after(cursor, text)
            inserted.append(p)
            cursor = p

    # Move the new paragraphs before the original next-heading paragraph.
    # We insert after the heading first because python-docx has no direct
    # insert-before API for paragraphs; then we rotate the XML nodes.
    for p in inserted:
        heading._p.addprevious(p._p)


def append_to_end(doc: Document, blocks: list[tuple[str, list[str]]]) -> None:
    cursor = doc.paragraphs[-1]
    for title, paragraphs in blocks:
        cursor = add_heading_after(cursor, title)
        for text in paragraphs:
            cursor = add_paragraph_after(cursor, text)


CHAPTER_BLOCKS = {
    "2 启动管理流程二": [
        (
            "1.3 源码实现范围补充说明",
            [
                "结合当前仓库源代码重新核对后，项目实际实现范围已经超过早期调研时的简单商品库存管理。系统代码按 app/routes、app/services、app/models、app/templates 和 app/static/js 分层组织，已经落地账号认证、商品管理、库存管理、收银结算、销售查询、财务管理、经营分析、公告管理、智能助手，以及会员、员工、供应商、系统参数等二期主数据模块。因此第 1 章中的项目目标应理解为一个覆盖门店经营闭环的管理系统，而不是只完成若干静态页面。",
                "从角色角度看，管理员不仅维护商品，还负责管理员注册审核、公告发布、财务对账、应付账款、月结、二期主数据和系统参数；收银员主要使用收银台、查看公告并完成销售结算；店长或财务人员关注经营分析、销售趋势、热销商品、日结差异和供应商应付。这个角色划分能够对应到代码中的 admin_required、cashier_required 和 role_required 等权限装饰器，说明调研结果已经转化成了可执行的访问控制。",
                "源码还体现了可演示性要求。应用启动时会初始化日志、数据库表、默认管理员、默认收银员、默认分类、二期演示数据，并在商品表为空时导入 seed_data.sql。这样项目从 Gitee 拉取后能够直接形成可演示数据库 data/supermarket.db，避免答辩时只有代码而没有业务数据。第 1 章需要把这一点写清楚，因为它解释了为什么系统可以现场运行、截图和测试。",
            ],
        )
    ],
    "3 需求建模": [
        (
            "2.4 源码结构与启动初始化补充说明",
            [
                "技术方案在源码中具体表现为 Flask 应用工厂结构。run.py 只负责启动应用，create_app 位于 app/__init__.py，负责加载配置、初始化 SQLAlchemy、注册路由、创建数据库表、初始化默认数据和设置请求日志。这样的结构比把全部逻辑写在一个文件中更清晰，便于把路由、服务、模型和测试分开维护。",
                "项目后端按三层理解：routes 目录负责页面入口和 API 入口，services 目录负责业务规则，models 目录负责 ORM 数据模型。以前端为例，templates 目录保存 Jinja2 页面模板，static/js 目录保存每个页面的异步交互脚本。商品、库存、收银、销售、财务、公告、分析、助手和二期模块都能在这几个目录中找到对应文件，说明技术方案不是停留在报告结构图，而是落实到了源码目录。",
                "启动初始化也是技术方案的一部分。系统会创建默认账号 admin/admin123 和 cashier01/123456，创建默认分类，并通过 seed_data.sql 填充演示商品、库存和交易数据。二期模块还会初始化会员、员工、供应商和系统参数示例。这样设计的目的，是保证报告截图、自动化测试、课堂演示和答辩口径使用同一套数据基础。",
                "智能助手的外部模型调用也属于技术方案扩展点。app/services/assistant.py 中预留了 LLM_PROVIDER、LLM_API_KEY、LLM_BASE_URL、LLM_MODEL 等配置，接口格式兼容常见 OpenAI 风格 chat completions。如果没有配置外部模型，系统会自动使用本地规则回复，因此演示不会依赖网络和第三方服务的稳定性。",
            ],
        )
    ],
    "4 数据建模": [
        (
            "3.4 源码功能与需求细化补充说明",
            [
                "需求建模需要进一步反映源码中已经实现的功能。智能助手不只是菜单项，而是由 /assistant 页面和 /api/assistant/chat 接口组成。前端通过 fetch 异步提交用户问题，后端先进行管理员权限校验，再解析 message 字段，空问题或超过 500 字符的问题会被拒绝。服务层优先匹配 FAQ 预设问题，然后根据库存、销售、商品和帮助类关键词调用已有业务服务生成回答；如果配置外部 LLM，则尝试调用大模型；如果失败，则降级成本地规则回复。",
                "商品批量导入也需要在需求层细化。代码中同时提供 import_products_from_csv 和 import_products_from_excel，说明用户既可以上传 CSV，也可以上传 Excel。导入过程不是简单读文件，而是检查必填字段、商品编码重复、条码重复、价格格式、库存数量和分类信息；合法记录会创建商品并初始化库存，不合法记录需要返回失败原因和统计结果。这一需求对应管理员快速建立商品主数据的真实工作场景。",
                "财务模块的需求不应只写成“查看财务”。源码中已经实现财务概览、收支流水、新增收支、日结对账、供应商应付、付款登记、月结快照和最近月结查询。日结对账需要按支付方式比较系统金额和实际金额，应付账款需要记录供应商、账单号、到期日和付款状态，月结快照需要固化收入、支出、毛利、净利润和现金流入流出。这些需求共同说明系统具备经营管理闭环。",
                "公告模块也不是普通文本列表。代码支持公告发布、上下线、目标角色、公告列表、未读数量、单条标记已读和全部标记已读。需求上应明确公告既有内容管理，也有阅读状态追踪。announcement_reads 表使系统能够判断某个用户是否读过某条公告，并支持未读提醒，这比在公告表中放一个简单状态字段更符合多用户场景。",
                "二期模块在需求上包括会员、员工、供应商和系统参数四组主数据。会员支持新增、编辑、搜索、启停和积分调整；员工支持工号、岗位、排班和状态维护；供应商支持供应商编码、联系人、电话、结算周期和启停；系统参数支持门店名称、库存预警开关、小票页脚等键值配置。每组模块都有正常路径和异常路径，例如重复手机号、重复工号、非法状态、积分扣减后为负数、系统参数键为空等。",
            ],
        )
    ],
    "5 界面设计": [
        (
            "4.3 源码数据结构补充说明",
            [
                "结合源码，数据建模需要把已实现模型写得更完整。账号与权限包括 users 和 admin_signup_requests，后者保存管理员注册申请、审核状态、审核人、审核时间和拒绝原因，使管理员权限不会通过普通注册直接获得。商品库存包括 categories、products、inventory 和 inventory_logs，商品主数据与当前库存、库存流水分开存储，既能支持查询，也能支持追溯。",
                "销售与财务数据采用业务单据和流水分离的思路。sales 保存销售单头，sale_items 保存成交时的商品、数量、单价和小计，避免后续商品调价影响历史订单。finance_transactions 保存收支流水，cash_reconciliations 保存日结对账，supplier_payables 与 payable_payments 保存供应商应付和付款记录，finance_period_closings 保存月结快照。这些表共同支撑财务概览、对账、应付和月结功能。",
                "公告模块使用 announcements 和 announcement_reads 两张表。announcements 负责标题、内容、级别、目标角色和发布状态；announcement_reads 负责记录用户对公告的阅读时间。这样一条公告可以被多个用户阅读，一个用户也可以阅读多条公告，数据库关系更清晰，也便于实现未读数量统计和全部已读功能。",
                "二期主数据模型包括 members、employees、suppliers 和 system_settings。members 保存手机号、等级、积分和状态，employees 保存员工编号、岗位、排班和状态，suppliers 保存供应商编码、联系人和结算周期，system_settings 使用键值对保存系统配置。键值对模型适合保存门店名称、库存预警开关、小票页脚等变化灵活但数量不大的配置项。",
                "智能助手本身没有单独建聊天记录表，当前实现采用即时问答模式。它通过 Product、Inventory、Category、Sale、SaleItem 等已有业务数据组织回答，并可调用库存汇总、低库存预警、销售概览和热销商品等服务函数。这样设计减少了额外表结构，同时体现助手模块与业务数据模型之间的依赖关系。",
            ],
        )
    ],
    "6 系统实施一": [
        (
            "5.3 源码界面交互补充说明",
            [
                "界面设计应补充智能助手页面。assistant.html 采用左侧快捷操作、右侧聊天窗口和底部输入框结构，assistant.js 负责字数统计、发送按钮状态、回车发送、清空对话、加载提示、用户消息和助手消息渲染。用户点击“查询库存”“销售分析”“低库存预警”“使用帮助”等快捷按钮时，前端会自动组织问题并调用 /api/assistant/chat。这个页面体现了系统的扩展创新点，不应只在报告中写一个模块名称。",
                "商品管理页面除了列表、搜索、分页和新增编辑弹窗，还包含导入交互。product.js 支持选择文件、拖拽上传、调用 /api/products/import，并把导入成功数、失败数和错误信息反馈给管理员。界面设计中应说明批量导入是为了减少商品主数据初始化工作量，和后端 CSV/Excel 解析逻辑相对应。",
                "财务页面不是单一表格，而是由财务概览、收支流水、日结对账、应付账款和月结操作共同组成。finance.js 会调用 overview、transactions、reconciliation、payables、closings 和 close-month 等接口，支持筛选、分页、保存对账金额、登记付款和生成月结快照。界面说明应强调财务人员可以从明细操作进入周期性汇总。",
                "公告交互需要体现未读状态。base.js 会在全局导航中请求 /api/announcements/unread-count，并支持公告弹窗、单条标记已读和全部标记已读；announcements.html 负责管理员发布和上下线公告。这样公告模块既有后台维护界面，也有前台提醒与阅读交互。",
                "二期主数据页面使用 master_data.html 与 master-data.js 复用一套界面逻辑。会员、员工、供应商和系统参数共用搜索、状态筛选、分页、弹窗编辑和启停操作，会员额外支持积分调整。这个复用设计减少了重复页面，也说明前端实现不是简单堆页面，而是抽象出统一的主数据管理交互模式。",
            ],
        )
    ],
    "7 系统实施二": [
        (
            "6.4 源码实现与测试映射补充说明",
            [
                "系统实施阶段需要把路由、服务和测试对应起来。认证模块由 auth.py 处理登录、注册、管理员申请审核和退出；common.py 提供 login_required、role_required、admin_required 和 cashier_required。测试中覆盖了登录、管理员审核、未登录访问和角色越权，说明权限边界不是只写在报告里，而是通过自动化测试验证。",
                "商品与库存模块的实现链路为 product 路由接收请求，products 服务完成商品创建、编辑、上下架、删除和 CSV/Excel 导入，inventory 服务完成库存设置、库存汇总、库存列表、库存流水和低库存预警。核心测试会验证商品创建后库存记录是否生成，库存调整后是否写入流水，导入 Excel 后是否同步生成商品和库存。",
                "收银和销售模块通过 cashier 与 sales 两组服务衔接。checkout_cashier_order 会校验商品状态、库存数量、支付方式和优惠金额，成功后生成销售单和销售明细，扣减库存并写入库存流水；get_sales_orders 和 get_sales_order_detail 支持订单列表、筛选和详情查询。测试需要覆盖库存不足失败、结算成功扣库存、订单详情金额正确等场景。",
                "财务模块通过 finance 服务实现财务概览、流水创建、日结对账、供应商应付、付款登记和月结快照。测试中保存对账记录后，需要验证系统金额、实际金额和差异金额计算正确。公告模块通过 announcements 服务实现公告列表、未读数量、标记已读和全部已读，测试会验证阅读后未读数减少。",
                "智能助手模块的实现链路为 assistant.html 与 assistant.js 提供聊天界面，routes/assistant.py 提供 /assistant 和 /api/assistant/chat，services/assistant.py 负责 FAQ、本地规则、库存查询、销售概览、商品帮助、大模型调用和降级回复。测试和演示时可选择库存预警、今日销售、如何导入商品等问题，证明助手不是装饰性页面，而是能调用系统业务数据的扩展模块。",
                "二期模块由 second_phase.py 提供会员、员工、供应商和系统参数服务。自动化测试覆盖正常 CRUD、启停、积分调整和异常路径，包括重复手机号、重复工号、非法状态、积分不足、参数键为空等。测试使用内存 SQLite，不污染 data/supermarket.db；coverage report 当前保留 100% 结果，用于证明后端模型和二期核心服务可验证。",
            ],
        )
    ],
}

END_BLOCKS = [
    (
        "7.4 源码覆盖与最终交付补充说明",
        [
            "最终交付时，报告需要与源码保持一致。当前项目源码中已实现的主要功能包括：账号登录与管理员审核、商品 CRUD 与 CSV/Excel 导入、库存汇总与流水、收银结算、销售订单查询、财务流水与日结对账、供应商应付与月结快照、经营分析、公告发布与已读、智能助手、会员管理、员工管理、供应商管理和系统参数维护。报告如果只写“商品库存销售管理”，会低估系统实际完成度。",
            "智能助手应作为实现亮点写入总结。该模块具备独立页面、聊天式交互、快捷问题、FAQ、库存问答、销售问答、商品问答、帮助问答、外部 LLM 预留接口和本地规则降级机制。它与库存、销售、商品服务联动，能够在没有外部模型的情况下稳定演示，在有模型配置时又能扩展到更自然的问答能力。",
            "测试交付应明确对应关系。tests/test_core_services.py 覆盖商品、库存、收银、销售详情、财务对账和公告已读；tests/test_auth_and_routes.py 覆盖登录、管理员审核、权限拦截和二期页面访问；tests/test_second_phase.py 覆盖会员、员工、供应商和系统参数正常路径；tests/test_second_phase_errors.py 覆盖二期异常路径；tests/test_api_functional.py 覆盖二期 API 功能流。pytest 通过结果和 coverage 100% 结果已经归档到 screenshots 目录。",
            "运行和数据交付也需要在总结中说明。README.md 提供 uv 环境、启动命令、默认账号、数据库位置和测试命令；data/supermarket.db 保存可演示数据库；data/SQL/schema.sql 和 seed_data.sql 保存结构与演示数据来源；supermarket-management-diagrams 保存 PNG 图表，supermarket-management-diagrams-drawio-editable 保存可编辑 drawio 源文件。这些文件共同构成 Gitee 提交物，而不是只有一个 Word 报告。",
            "因此，第 7 章的结论应强调：本项目已经形成从需求、数据模型、界面、后端服务、测试、截图、PPT 到 Gitee 仓库归档的完整交付链。后续答辩时可以按“业务需求 -> 代码模块 -> 数据表 -> 测试用例 -> 报告图表”的顺序回答追问，尤其要主动说明智能助手、财务月结、商品导入、公告已读和二期主数据这些容易被忽略但源码已经实现的功能。",
        ],
    )
]


def main() -> None:
    report = target_report()
    backup = report.with_name(report.stem + "_before_chapter_source_expansion.docx")
    if not backup.exists():
        shutil.copy2(report, backup)

    doc = Document(report)
    # Insert from back to front so heading lookup remains stable.
    for heading_text, blocks in reversed(list(CHAPTER_BLOCKS.items())):
        insert_before_heading(doc, heading_text, blocks)
        print(f"expanded before {heading_text}")

    append_to_end(doc, END_BLOCKS)
    doc.save(report)
    print("saved", report)
    print("backup", backup)
    print("paragraphs", len(Document(report).paragraphs))


if __name__ == "__main__":
    main()
