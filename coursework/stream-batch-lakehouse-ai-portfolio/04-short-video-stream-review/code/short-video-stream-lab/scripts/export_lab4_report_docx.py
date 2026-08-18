"""Export the Lab 4 Markdown report to a Word document with Microsoft YaHei."""

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
LAB_ROOT = ROOT.parents[1]
SOURCE = LAB_ROOT / "短视频审核实验报告_[REDACTED]_Bohan Yu.md"
TARGET = LAB_ROOT / "短视频审核实验报告_[REDACTED]_Bohan Yu.docx"


def _set_font(run, size: int = 10) -> None:
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def _set_paragraph_font(paragraph, size: int = 10) -> None:
    for run in paragraph.runs:
        _set_font(run, size)


def _set_cell_font(cell, size: int = 9) -> None:
    for paragraph in cell.paragraphs:
        _set_paragraph_font(paragraph, size)


def _parse_table(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in lines:
        stripped = line.strip().strip("|")
        if set(stripped.replace("|", "").replace("-", "").replace(" ", "")) == set():
            continue
        rows.append([cell.strip().strip("`") for cell in stripped.split("|")])
    return rows


def _add_table(document: Document, lines: list[str]) -> None:
    rows = _parse_table(lines)
    if not rows:
        return
    table = document.add_table(rows=len(rows), cols=max(len(row) for row in rows))
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            cell = table.cell(i, j)
            cell.text = value
            _set_cell_font(cell)


def _add_code_block(document: Document, lines: list[str]) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = document.styles["Normal"]
    run = paragraph.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F3F4F6")
    paragraph._p.get_or_add_pPr().append(shading)


def _add_markdown_line(document: Document, line: str) -> None:
    stripped = line.strip()
    if not stripped:
        return
    if stripped.startswith("# "):
        paragraph = document.add_heading(stripped[2:], level=1)
        _set_paragraph_font(paragraph, 18)
    elif stripped.startswith("## "):
        paragraph = document.add_heading(stripped[3:], level=2)
        _set_paragraph_font(paragraph, 15)
    elif stripped.startswith("### "):
        paragraph = document.add_heading(stripped[4:], level=3)
        _set_paragraph_font(paragraph, 13)
    elif stripped.startswith("- "):
        paragraph = document.add_paragraph(stripped[2:], style="List Bullet")
        _set_paragraph_font(paragraph, 10)
    elif len(stripped) > 2 and stripped[0].isdigit() and stripped[1] == ".":
        paragraph = document.add_paragraph(stripped, style="List Number")
        _set_paragraph_font(paragraph, 10)
    else:
        paragraph = document.add_paragraph(stripped.replace("`", ""))
        _set_paragraph_font(paragraph, 10)


def main() -> None:
    """Create a docx report from the Markdown source."""
    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if stripped.startswith("```"):
            block: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            _add_code_block(document, block)
        elif stripped.startswith("|"):
            table_lines: list[str] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            _add_table(document, table_lines)
            continue
        else:
            _add_markdown_line(document, line)
        i += 1

    document.save(TARGET)
    print(TARGET)


if __name__ == "__main__":
    main()
