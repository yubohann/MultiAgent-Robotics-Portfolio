"""Update the Lab 4 Word report with screenshots while keeping its cover."""

from __future__ import annotations

from pathlib import Path
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
LAB_ROOT = PROJECT_ROOT.parents[1]
PORTFOLIO_ROOT = PROJECT_ROOT.parents[2]
SCREENSHOT_DIR = LAB_ROOT / "screenshots_[REDACTED]"
REPORT = next(
    p
    for p in LAB_ROOT.glob("*.docx")
    if "[REDACTED]" in p.name and "before" not in p.name and "backup" not in p.name
)
BACKUP = LAB_ROOT / "短视频审核实验报告_[REDACTED]_Bohan Yu_before_screenshot_update.docx"
STRENGTHEN_BACKUP = LAB_ROOT / "短视频审核实验报告_[REDACTED]_Bohan Yu_before_strengthen.docx"


SCREENSHOTS = {
    "source": ["00_public_drone_source_license_[REDACTED].png"],
    "env": [
        "01_env_python_ffmpeg_ollama_[REDACTED].png",
        "02_python_dependencies_[REDACTED].png",
    ],
    "model": [
        "03_ollama_8b_vision_model_1_[REDACTED].png",
        "03_ollama_8b_vision_model_2_[REDACTED].png",
        "04b1_baseline_comparison_[REDACTED].png",
        "11_model_selector_all_models_[REDACTED].png",
        "11_1_model_selector_all_models_[REDACTED].png",
    ],
    "run": [
        "04_verify_demo_8b_vlm_passed_[REDACTED].png",
        "05_fastapi_server_start_[REDACTED].png",
    ],
    "api": [
        "06_upload_processing_visible_[REDACTED].png",
        "07_api_videos_[REDACTED].png",
        "08_api_health_jobs_[REDACTED].png",
        "09_api_events_trace_[REDACTED].png",
        "10_final_review_result_[REDACTED].png",
        "10_1_final_review_result_[REDACTED].png.png",
    ],
    "kafka": [
        "12_kafka_topics_[REDACTED].png",
        "13_kafka_ingest_and_result_[REDACTED].png",
    ],
}


CAPTIONS = {
    "00_public_drone_source_license_[REDACTED].png": "图 4-1 Pexels 公开视频来源与 License 证据。",
    "01_env_python_ffmpeg_ollama_[REDACTED].png": "图 4-2 Python、FFmpeg、FFprobe、Ollama 环境验证。",
    "02_python_dependencies_[REDACTED].png": "图 4-3 Python 虚拟环境依赖安装与导入验证。",
    "03_ollama_8b_vision_model_1_[REDACTED].png": "图 4-4 Ollama 模型列表与本地 8B Vision 主模型验证。",
    "03_ollama_8b_vision_model_2_[REDACTED].png": "图 4-5 Qwen3-VL 4B 与 Gemma 3 4B 两个视觉基线验证。",
    "04b1_baseline_comparison_[REDACTED].png": "图 4-6 本机主模型、两个 4B 基线与 OpenCV 规则基线对照结果。",
    "04_verify_demo_8b_vlm_passed_[REDACTED].png": "图 4-7 严格模式 verify_demo.py 端到端验收通过。",
    "05_fastapi_server_start_[REDACTED].png": "图 4-8 FastAPI 服务启动并监听 127.0.0.1:5050。",
    "06_upload_processing_visible_[REDACTED].png": "图 4-9 自选无人机视频上传后立即显示 processing 卡片。",
    "07_api_videos_[REDACTED].png": "图 4-10 /api/videos 展示视频状态、标签、风险分和模型后端。",
    "08_api_health_jobs_[REDACTED].png": "图 4-11 /api/health 展示 worker、jobs 与 active_model 状态。",
    "09_api_events_trace_[REDACTED].png": "图 4-12 /api/events 展示 ingest、queued、worker、frame_sample、vlm_understanding 等事件链路。",
    "10_final_review_result_[REDACTED].png": "图 4-13 网页最终审核结果：摘要、标签、风险分和模型信息。",
    "10_1_final_review_result_[REDACTED].png.png": "图 4-14 网页最终结果补充视图。",
    "11_model_selector_all_models_[REDACTED].png": "图 4-15 网页模型选择区展示主模型、两个主流视觉基线、规则基线与备用模型。",
    "11_1_model_selector_all_models_[REDACTED].png": "图 4-16 命令行模型清单展示 downloaded、active、backend 与 memory_tier。",
    "12_kafka_topics_[REDACTED].png": "图 4-17 Kafka 学号 topic 创建结果。",
    "13_kafka_ingest_and_result_[REDACTED].png": "图 4-18 Kafka producer 发送视频进入事件并从 result topic 获得审核结果。",
}


def set_run_font(run, size: int = 10, bold: bool | None = None) -> None:
    """Apply Microsoft YaHei to one run."""
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    if bold is not None:
        run.bold = bold


def set_paragraph_font(paragraph, size: int = 10) -> None:
    """Apply font to every run in a paragraph."""
    for run in paragraph.runs:
        set_run_font(run, size)


def add_heading(document: Document, text: str, level: int = 2) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = f"Heading {level}"
    run = paragraph.add_run(text)
    set_run_font(run, 16 if level == 2 else 13, bold=True)


def add_body(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Pt(20)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(text)
    set_run_font(run, 10)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        run = paragraph.add_run(item)
        set_run_font(run, 10)


def add_code(document: Document, lines: list[str]) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F3F4F6")
    paragraph._p.get_or_add_pPr().append(shading)


def add_table(document: Document, rows: list[list[str]]) -> None:
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            cell = table.cell(i, j)
            cell.text = value
            for paragraph in cell.paragraphs:
                set_paragraph_font(paragraph, 9)
                if i == 0:
                    for run in paragraph.runs:
                        run.bold = True


def add_screenshot(document: Document, filename: str) -> None:
    path = SCREENSHOT_DIR / filename
    if not path.exists():
        add_body(document, f"截图文件缺失：{filename}")
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(5.85))
    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption.add_run(CAPTIONS.get(filename, filename))
    set_run_font(caption_run, 9)
    caption_run.italic = True


def clear_after_cover(document: Document, cover_paragraphs: int = 10) -> None:
    """Keep the first cover paragraphs and remove the original body."""
    body = document._body._element
    kept = 0
    for element in list(body):
        if element.tag.endswith("}sectPr"):
            continue
        if element.tag.endswith("}p") and kept < cover_paragraphs:
            kept += 1
            continue
        body.remove(element)


def set_document_defaults(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    for section in document.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)


def build_report(document: Document) -> None:
    """Append the updated report body after the preserved cover."""
    document.add_page_break()

    add_heading(document, "一、实验目的")
    add_body(
        document,
        "本综合实验围绕实时短视频流的智能审核与发布展开，目标是构建一个能够完成视频上传、异步入队、关键帧抽样、多模态理解、风险审核、标签生成、发布展示和事件追踪的完整闭环。"
        "实验口径与前三个实验保持一致：用可运行系统、命令行日志、网页截图和事件链路共同证明结果，而不是只给出静态说明。"
    )
    add_body(
        document,
        "从业务角度看，短视频平台的审核难点不在于“能否上传文件”，而在于上传后如何在用户体验、审核准确性和系统吞吐之间取得平衡。"
        "如果上传接口同步等待多模态模型推理，用户会在网络请求中长时间等待，模型超时还会导致上传失败；如果完全放弃模型理解，又无法解释视频主题、风险线索和推荐标签。"
        "因此，本实验采用“上传立即可见 + 后台异步审核 + 事件可追踪”的设计，将用户交互链路和模型推理链路解耦。"
    )
    add_bullets(
        document,
        [
            "掌握 FastAPI 上传接口与前端 processing 即时可见的异步交互方式。",
            "掌握 OpenCV/FFmpeg 对短视频元数据、亮度、运动、闪烁和关键帧的处理方法。",
            "使用本地 Ollama 多模态视觉模型完成摘要、标签和风险线索识别。",
            "使用 SQLite jobs/events 模拟工业消息队列和可观测日志，使用 Kafka topic 完成加分流式扩展。",
            "通过截图、API 输出和日志形成带学号 [REDACTED] 的完整实验追踪证据链。",
        ],
    )
    add_table(
        document,
        [
            ["实验问题", "本实验解决方式", "证据位置"],
            ["上传后等待模型导致体验差", "先写 processing 记录并返回页面，worker 后台完成审核", "图 4-9、图 4-12"],
            ["模型输出不可追溯", "保留 /api/videos、/api/events 和日志文件", "图 4-10、图 4-12"],
            ["单模型结论缺少对照", "部署 8B 主模型、两个 4B 视觉基线和 OpenCV 规则基线", "图 4-4 至图 4-6、图 4-15 至图 4-16"],
            ["单机 Demo 与工业链路脱节", "使用 SQLite 模拟本地队列，并补充 Kafka topic 加分链路", "图 4-17、图 4-18"],
        ],
    )

    add_heading(document, "二、实验环境与公开素材")
    add_body(
        document,
        "实验运行环境为 Windows PowerShell，本机 GPU 为 NVIDIA GeForce RTX 4090。项目不依赖全局 Python 或全局 FFmpeg，而是使用项目内虚拟环境和便携 FFmpeg，"
        "以避免 Windows 中文路径、PATH 冲突和 OpenCV 写图失败等问题。"
    )
    add_table(
        document,
        [
            ["项目", "取值"],
            ["姓名 / 学号", "Bohan Yu / [REDACTED]"],
            ["班级", "23计算师资1班"],
            ["运行终端", "Windows PowerShell"],
            ["GPU", "NVIDIA GeForce RTX 4090"],
            ["项目目录", "code/short-video-stream-lab"],
            ["截图目录", "screenshots_[REDACTED]"],
            ["日志目录", "logs_[REDACTED]"],
        ],
    )
    add_body(
        document,
        "本实验自选素材来自 Pexels 公开视频，报告中保留视频 URL 和 License URL。Pexels 公开视频可免费使用，署名不是必需项，但本报告仍保留来源链接作为可追溯证据。"
    )
    add_body(
        document,
        "选择公开视频而不是随意上传互联网视频，是为了同时满足版权合规和内容安全。实验中的无人机航拍、夜间街景、舞台灯光和运动场跑步都来自公开授权素材，"
        "其中无人机航拍用于主业务场景，低照度与强闪烁素材用于验证复核规则的区分度，运动场素材用于证明正常公开视频可以自动发布。"
    )
    add_code(
        document,
        [
            "Pexels public video",
            "https://videos.pexels.com/video-files/7457249/7457249-uhd_3840_2160_25fps.mp4",
            "License: https://www.pexels.com/license/",
            "实验压缩版: data/incoming/upload_[REDACTED]_yubohan_drone_review.mp4",
        ],
    )
    for filename in SCREENSHOTS["source"] + SCREENSHOTS["env"]:
        add_screenshot(document, filename)

    add_heading(document, "三、系统总体架构")
    add_body(
        document,
        "系统由前端页面、FastAPI API 服务、本地媒体目录、SQLite jobs/events、后台 worker、OpenCV/FFmpeg 预处理模块和 Ollama VLM 模型服务组成。"
        "上传请求只负责保存视频、写入 processing 记录和创建后台任务；耗时的视频理解与审核由 worker 异步完成。"
    )
    add_body(
        document,
        "这个架构对应真实短视频平台中的常见分层：入口服务负责接收视频并生成初始状态，消息队列负责削峰和异步调度，媒体存储保存视频与关键帧，"
        "模型服务负责视觉语义理解，审核规则负责把模型输出和可解释指标转成 published、review 或 rejected。"
        "本实验虽然使用单机 SQLite 和本地目录实现，但接口边界与工业系统保持一致，后续可以替换成 Kafka、MinIO 和分布式模型服务。"
    )
    add_table(
        document,
        [
            ["模块", "职责", "工业化映射"],
            ["FastAPI", "上传、查询、模型选择、静态媒体服务", "API Gateway / Backend Service"],
            ["SQLite jobs/events", "本地队列、任务状态和事件追踪", "Kafka / Pulsar / PostgreSQL"],
            ["data/media", "保存上传视频、转码文件、缩略图和关键帧", "MinIO / S3 / OSS"],
            ["OpenCV/FFmpeg", "抽帧、亮度、运动、闪烁和视频元数据", "视频预处理服务"],
            ["Ollama VLM", "多模态理解、摘要、标签和风险线索识别", "vLLM / SGLang / 模型服务集群"],
        ],
    )

    add_heading(document, "四、本地多模态模型与基线对照")
    add_body(
        document,
        "主模型选择本机已安装并验证具备 vision capability 的 ministral-3:8b，项目内部 ID 为 ministral-3-8b-ollama。"
        "为体现模型对照工作量，同时避免硬盘继续拉取大型模型，本实验新增并部署 qwen3-vl:4b 与 gemma3:4b 两个主流 4B 视觉基线，并保留 OpenCV Local Baseline 作为无模型权重规则兜底。"
    )
    add_table(
        document,
        [
            ["角色", "项目模型 ID", "Ollama 标签", "用途"],
            ["主模型", "ministral-3-8b-ollama", "ministral-3:8b", "最高分主链路证据"],
            ["视觉基线 1", "qwen3-vl-4b-ollama", "qwen3-vl:4b", "Qwen 系列主流视觉基线"],
            ["视觉基线 2", "gemma3-4b-ollama", "gemma3:4b", "Gemma 系列主流视觉基线"],
            ["规则基线", "local-baseline", "no-weight", "OpenCV 亮度/运动/闪烁规则兜底"],
            ["轻量备用", "ministral-3-3b-ollama", "ministral-3:3b", "已安装轻量备用模型"],
        ],
    )
    add_body(
        document,
        "对照脚本覆盖正常无人机航拍、Pexels 真实夜间低照度街景和 Pexels 真实舞台灯光闪烁片段三个场景，"
        "因此能够看到 published/review 以及 0.0、42.0、48.0 三档风险分数的区分。"
    )
    add_table(
        document,
        [
            ["场景", "真实素材", "预期状态", "风险分数", "审核含义"],
            ["normal_drone", "Pexels 无人机航拍", "published", "0.0", "常规航拍/巡检素材，未命中高风险信号"],
            ["low_light_review", "Pexels 夜间低照度街景", "review", "42.0", "画面过暗，自动理解置信度下降，需要人工复核"],
            ["flash_risk_review", "Pexels 舞台灯光闪烁", "review", "48.0", "强亮度跳变和高运动幅度可能造成观看不适"],
        ],
    )
    add_body(
        document,
        "这里的对照不只是比较模型速度，还比较模型能否给出业务语义。OpenCV baseline 能稳定计算亮度、运动和闪烁，因此适合作为兜底；"
        "但它不能真正理解“无人机巡检、居民区、舞台灯光、夜间街景”等语义。Qwen3-VL 和 Gemma 3 作为 4B 级视觉基线，证明同一套 pipeline 可以切换不同 VLM；"
        "Ministral 3 8B 则作为本机主模型提供更强的语义描述和最高分主链路证据。"
    )
    for filename in SCREENSHOTS["model"]:
        add_screenshot(document, filename)

    add_heading(document, "五、端到端运行与网页证据")
    add_body(
        document,
        "严格验收时设置 ALLOW_LOCAL_MODEL_FALLBACK=0，表示如果本地多模态模型不可用，脚本必须失败，不能悄悄降级到规则 baseline。"
        "verify_demo.py 运行三段真实公开视频样例，最终得到 1 个 published、2 个 review、0 个 rejected，证明主链路能够稳定运行。"
    )
    add_code(
        document,
        [
            "$env:STUDENT_ID = \"[REDACTED]\"",
            "$env:ALLOW_LOCAL_MODEL_FALLBACK = \"0\"",
            ".\\.venv\\Scripts\\python.exe scripts\\verify_demo.py",
        ],
    )
    add_body(
        document,
        "随后启动 FastAPI 服务并访问 http://127.0.0.1:5050。上传自选无人机视频后，页面立即出现 processing 卡片；后台 worker 完成抽帧和模型理解后，页面再更新为最终审核结果。"
    )
    add_body(
        document,
        "图 4-9 的关键意义是证明系统没有把“模型推理完成”作为上传成功的前置条件。视频文件和标题先进入页面，摘要、标签和风险分数可以稍后由 worker 补齐。"
        "这种设计适合短视频平台的实际场景：用户侧关注上传是否成功，平台侧则通过后台审核逐步补充理解结果和发布状态。"
    )
    for filename in SCREENSHOTS["run"] + SCREENSHOTS["api"]:
        add_screenshot(document, filename)

    add_heading(document, "六、API 与事件链路分析")
    add_body(
        document,
        "/api/videos 展示视频标题、状态、标签、风险分、模型 selected_id 和 backend；/api/health 展示 jobs queued/running/done、worker 与 active_model；"
        "/api/events 展示 ingest、stream、queued、worker、frame_sample、preprocess、vlm_understanding、moderation、publish 等阶段。"
    )
    add_body(
        document,
        "最终无人机视频被判定为 published，risk_score 为 0.0。模型摘要识别到居民区鸟瞰、道路、房屋、绿化带等内容，未命中本实验策略中的高风险信号。"
        "虽然本次自动发布，但无人机画面仍然天然适合审核：它可能暴露地理位置、住宅、车牌、人脸、厂区、学校、交通节点或安防区域，因此需要多模态理解与规则审核。"
    )
    add_table(
        document,
        [
            ["API / 事件", "证明内容", "评分价值"],
            ["/api/videos", "视频标题、状态、摘要、标签、risk_score、selected_id、backend", "证明模型结果真实写入系统"],
            ["/api/health", "worker、jobs、active_model、ollama_ready", "证明后台队列和模型状态可观测"],
            ["/api/events", "ingest、queued、worker、frame_sample、vlm_understanding、moderation、publish", "证明审核链路可追踪"],
            ["网页最终结果", "视频播放器、摘要、标签、状态和模型后端同屏展示", "证明用户侧可见闭环完成"],
        ],
    )
    add_body(
        document,
        "无人机视频的审核结论不能简单理解为“没有风险所以不需要审核”。恰恰相反，无人机视频适合作为审核测试素材，是因为它具有高视角和广覆盖特征，"
        "容易把道路、住宅、车辆、基础设施、园区边界等信息同时收入画面。本次素材最终自动发布，说明当前画面未命中高风险规则；但平台仍需要通过关键帧抽样、多模态理解和事件日志，"
        "保留可解释依据，避免后续出现隐私、地理安全或飞行合规争议时无法追踪。"
    )

    add_heading(document, "七、Kafka 加分项")
    add_body(
        document,
        "在主链路完成后，本实验继续验证 Kafka 扩展路径。Windows 侧通过 WSL Docker 中的 bigdata-kafka 容器连接 localhost:9092，"
        "创建带学号的 short_video_ingest_[REDACTED] 和 short_video_result_[REDACTED] 两个 topic。"
    )
    add_body(
        document,
        "Kafka producer 发送一条无人机短视频进入事件，kafka_ai_consumer 调用同一套本地多模态审核链路处理后，把包含 status、risk_score 和 tags 的结果写入 result topic。"
        "其他低照度、舞台闪烁和运动场视频已在 FastAPI/SQLite 主链路与 baseline 对照中验证，Kafka 部分用于证明流式扩展能力。"
    )
    add_body(
        document,
        "Kafka 只选取无人机视频发送一条事件，并不是缺少其他样例，而是为了突出“流式事件进入 -> AI 审核消费者处理 -> result topic 输出”的链路验证。"
        "低照度和强闪烁样例已经在本机 baseline 对照中证明复核规则有效；Kafka 加分项的核心是证明该审核逻辑可以从本地 SQLite 队列平滑迁移到课程前三个实验使用的流处理思想。"
    )
    add_code(
        document,
        [
            "short_video_ingest_[REDACTED]",
            "short_video_result_[REDACTED]",
            "[REDACTED]_Bohan Yu_Kafka_无人机巡检短视频",
            "status: published",
            "risk_score: 0.0",
        ],
    )
    for filename in SCREENSHOTS["kafka"]:
        add_screenshot(document, filename)

    add_heading(document, "八、综合分析")
    add_body(
        document,
        "本实验的关键价值在于把短视频审核拆成可观测的异步链路：上传接口不等待模型完成，页面先展示 processing；"
        "后台 worker 负责关键帧抽样、多模态理解、风险评分和状态更新；events 日志解释每个视频为什么通过、复核或拒绝。"
    )
    add_bullets(
        document,
        [
            "Ministral 3 8B Vision 能够输出贴近业务的无人机巡检、居民区、基础设施检查等语义标签。",
            "Qwen3-VL 4B 与 Gemma 3 4B 作为主流视觉基线，证明系统并非只绑定一个模型，具备横向对照能力。",
            "OpenCV Local Baseline 运行最快，但只能给出亮度、运动、横竖屏等低层信号，适合作为兜底而不是主模型。",
            "低照度真实街景进入 review，risk_score=42.0；强闪烁真实舞台片段进入 review，risk_score=48.0；正常无人机视频自动发布，risk_score=0.0。",
            "Kafka topic 与 producer/consumer/result 证据说明该系统可以从单机 SQLite 队列扩展到课程前三个实验中的流式处理链路。",
        ],
    )
    add_body(
        document,
        "与前三个实验相比，本实验的重点从“数据流处理与存储治理”进一步扩展到“非结构化媒体 + 多模态模型 + 审核业务规则”。"
        "实验一二中的 Kafka、Flink、Paimon 关注数据在流式计算和数据湖中的吞吐、倾斜、小文件等问题；实验四则把 Kafka 事件思想用于视频进入事件，"
        "把模型服务放到异步 worker 后面，并把每一步处理结果写入可查询 API 和事件日志。两者共同体现了云计算与大数据处理课程中的核心能力：高吞吐进入、异步处理、状态可观测、结果可追溯。"
    )
    add_body(
        document,
        "本实验仍存在局限。第一，关键帧抽样不能覆盖视频中的每一帧，极短暂的人脸、车牌或敏感标识可能被漏掉；第二，VLM 对远距离小目标的识别能力有限，"
        "对地理位置和禁飞区的判断需要接入地图、GPS 或地理围栏；第三，当前规则分数是教学版策略，真实平台还需要分级策略、人工复核队列、灰度发布和申诉机制。"
        "这些局限也说明，多模态模型适合提供理解能力，但不能替代完整的内容安全工程体系。"
    )
    add_table(
        document,
        [
            ["潜在风险", "当前实验处理", "生产环境改进"],
            ["远距离人脸/车牌漏检", "VLM 关键帧理解 + 规则复核", "接入目标检测、OCR、车牌/人脸模糊化"],
            ["敏感地理位置暴露", "人工复核原因与事件日志保留", "结合 GPS、地图围栏和敏感区域库"],
            ["模型误判或输出不稳定", "多模型基线对照 + local-baseline 兜底", "模型评测集、A/B 测试、人工抽检"],
            ["单机服务吞吐有限", "SQLite jobs 演示异步队列", "Kafka/Pulsar + 多 worker + 模型服务池"],
        ],
    )

    add_heading(document, "九、实验结论")
    add_body(
        document,
        "本次综合实验完成了实时短视频上传、异步队列、多模态理解、风险审核、自动标签、发布展示、API 查询、事件追踪和 Kafka 扩展。"
        "所有核心证据均带有学号 [REDACTED]，素材来源、模型部署、网页效果、API 输出、日志和 Kafka topic 之间口径一致。"
    )
    add_body(
        document,
        "若部署到真实平台，本实验中的本地 data/media 应替换为 MinIO/S3，SQLite jobs/events 应替换为 Kafka/Pulsar 和 PostgreSQL/ClickHouse，"
        "Ollama 应替换为可扩缩容的 vLLM/SGLang 模型服务，同时接入 OCR、人脸/车牌检测、地理围栏和人工复核系统，以降低无人机视频中的隐私、安全和合规风险。"
    )
    add_body(
        document,
        "综合来看，本实验不仅完成了一个可运行的短视频审核 Demo，也展示了从单机验证到工业架构迁移的思路：先用本地可控环境保证功能闭环，再用模型基线和 Kafka 证据证明系统具备扩展空间。"
        "这种实现路径兼顾了课程实验的可复现性和真实业务系统的工程合理性。"
    )
    add_body(document, "本次实验追踪学号：[REDACTED]。")


def main() -> None:
    if not BACKUP.exists():
        shutil.copy2(REPORT, BACKUP)
    if not STRENGTHEN_BACKUP.exists():
        shutil.copy2(REPORT, STRENGTHEN_BACKUP)
    # Rebuild from the pre-screenshot copy so repeated runs do not preserve
    # a previously generated body as part of the cover.
    source = BACKUP if BACKUP.exists() else REPORT
    document = Document(str(source))
    set_document_defaults(document)
    clear_after_cover(document, cover_paragraphs=7)
    build_report(document)
    document.save(str(REPORT))
    print(REPORT)
    print(f"backup: {BACKUP}")


if __name__ == "__main__":
    main()
