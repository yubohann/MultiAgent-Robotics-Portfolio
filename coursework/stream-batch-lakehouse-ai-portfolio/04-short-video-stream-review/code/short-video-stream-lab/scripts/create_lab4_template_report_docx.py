"""Create a Lab 4 report docx that follows the earlier course-report format."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
PORTFOLIO_ROOT = ROOT.parents[2]
LAB_ROOT = ROOT.parents[1]
REFERENCE = PORTFOLIO_ROOT / "云计算与大数据处理_实验报告_实验一二_基于备份完成版.docx"
TARGET = LAB_ROOT / "云计算与大数据处理_实验报告_实验四_短视频审核_待插图版.docx"


def set_run_font(run, size: float = 11, bold: bool | None = None) -> None:
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    if bold is not None:
        run.bold = bold


def set_paragraph_font(paragraph, size: float = 11) -> None:
    for run in paragraph.runs:
        set_run_font(run, size)


def add_center(document: Document, text: str, size: float, bold: bool = False, space_after: int = 6) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_run_font(run, size, bold)
    paragraph.paragraph_format.space_after = Pt(space_after)


def add_heading_like(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run, 12, True)
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)


def add_body(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Pt(22)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(text)
    set_run_font(run, 10.5)


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Paragraph")
    run = paragraph.add_run("• " + text)
    set_run_font(run, 10.5)


def add_code(document: Document, lines: list[str]) -> None:
    for line in lines:
        paragraph = document.add_paragraph(style="内容块-16-a")
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def add_table(document: Document, rows: list[list[str]]) -> None:
    table = document.add_table(rows=len(rows), cols=max(len(row) for row in rows))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            cell = table.cell(i, j)
            cell.text = value
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, 10, i == 0)


def add_figure_placeholder(document: Document, caption: str) -> None:
    box = document.add_paragraph()
    box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = box.add_run("【此处插入截图】")
    set_run_font(run, 10.5, True)
    caption_p = document.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_p.add_run(caption)
    set_run_font(caption_run, 10)


def main() -> None:
    document = Document(REFERENCE)

    # Clear template body while preserving styles from the reference report.
    body = document._body._element
    for child in list(body):
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)

    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    add_center(document, "《云计算与大数据处理》", 18, True, 10)
    add_center(document, "实验报告", 22, True, 18)
    add_center(document, "实时短视频流智能审核与发布综合实验", 15, True, 10)

    add_table(
        document,
        [
            ["实验名称", "实时短视频流智能审核与发布综合实验"],
            ["实验日期", f"{date.today().year}年{date.today().month}月{date.today().day}日"],
            ["专业", "计算机科学与技术"],
            ["班级", "23计算师资1班"],
            ["姓名", "Bohan Yu"],
            ["学号", "[REDACTED]"],
            ["运行环境", "Windows PowerShell + FastAPI + SQLite + Ollama + OpenCV/FFmpeg"],
            ["GPU", "NVIDIA GeForce RTX 4090"],
            ["项目目录", str(LAB_ROOT)],
        ],
    )

    add_heading_like(document, "一、实验目的")
    for item in [
        "掌握短视频上传后立即可见、后台异步审核的系统设计方法。",
        "理解短视频平台中视频预处理、关键帧抽样、多模态理解、审核规则和自动发布之间的关系。",
        "使用本地 Ollama 多模态模型对无人机短视频进行摘要生成、标签生成和风险识别。",
        "通过 FastAPI、SQLite jobs/events、本地 worker 实现可观测的审核发布 Demo。",
        "将本地轻量实现映射到 Kafka、MinIO/S3、PostgreSQL/ClickHouse、模型服务等工业组件。",
    ]:
        add_bullet(document, item)

    add_heading_like(document, "二、实验环境与数据对象")
    add_body(
        document,
        "本实验在 Windows PowerShell 中运行。项目没有依赖全局 Python 和全局 FFmpeg，而是使用项目内 .venv、"
        "项目内便携 FFmpeg/FFprobe 和本机 Ollama 服务，避免 Windows 中文路径和系统 PATH 不一致造成环境问题。"
    )
    add_table(
        document,
        [
            ["类别", "配置或对象"],
            ["Python 环境", ".venv\\Scripts\\python.exe，Python 3.11.15"],
            ["视频工具", "tools\\ffmpeg\\bin\\ffmpeg.exe / ffprobe.exe，版本 8.1.1"],
            ["GPU", "NVIDIA GeForce RTX 4090"],
            ["Web 服务", "FastAPI，访问地址 http://127.0.0.1:5050"],
            ["异步队列", "SQLite jobs 表 + 本地 worker"],
            ["事件追踪", "SQLite events 表，API 为 /api/events"],
            ["主模型", "ministral-3-8b-ollama -> Ollama ministral-3:8b"],
            ["视觉基线", "qwen3-vl:4b、gemma3:4b、local-baseline"],
            ["自选视频", "upload_[REDACTED]_yubohan_drone_review.mp4"],
        ],
    )
    add_figure_placeholder(document, "图 4-1 Python、FFmpeg、FFprobe 与 Ollama 环境验证截图。")
    add_figure_placeholder(document, "图 4-2 Python 依赖与项目内 FFmpeg 路径验证截图。")

    add_heading_like(document, "三、总体架构与实验流程")
    add_body(
        document,
        "系统采用 FastAPI 提供上传接口和 Demo 网站。用户上传短视频后，服务端先保存媒体文件和 processing 状态，"
        "然后将审核任务写入 SQLite jobs 本地耐久队列。后台 worker 异步执行视频预处理、关键帧抽样、"
        "本地多模态模型理解、风险规则判定和发布状态更新。前端通过 API 轮询获取最新状态，"
        "因此视频可以先展示，摘要、标签和审核结果随后更新。"
    )
    add_table(
        document,
        [
            ["本实验模块", "作用", "工业系统可替换方案"],
            ["data/media 本地目录", "保存上传视频和缩略图", "MinIO / S3 / OSS"],
            ["SQLite jobs", "模拟异步任务队列", "Kafka / Pulsar / RabbitMQ"],
            ["SQLite videos/events", "保存元数据和事件日志", "PostgreSQL / ClickHouse / Paimon"],
            ["本地 worker", "执行视频理解和审核任务", "容器化消费者 / Kubernetes Job"],
            ["Ollama 本地模型", "执行多模态视觉理解", "vLLM / SGLang / Triton / 云模型服务"],
        ],
    )

    add_heading_like(document, "四、本地多模态模型选择")
    add_body(
        document,
        "本实验默认主模型为本机已安装的 ministral-3:8b，项目模型 ID 为 ministral-3-8b-ollama。"
        "ollama show 结果显示该模型参数规模为 8.9B，量化方式为 Q4_K_M，capabilities 包含 completion、vision、tools，"
        "能够接收关键帧图片完成结构化视频理解。为了满足模型对照要求，同时控制硬盘占用，本实验只新增两个约 3.3GB 的 4B 级视觉模型：qwen3-vl:4b 和 gemma3:4b。"
    )
    add_table(
        document,
        [
            ["角色", "项目模型 ID", "Ollama 标签", "用途"],
            ["主模型", "ministral-3-8b-ollama", "ministral-3:8b", "最高分主链路证据"],
            ["基线 1", "qwen3-vl-4b-ollama", "qwen3-vl:4b", "Qwen 系列主流视觉基线"],
            ["基线 2", "gemma3-4b-ollama", "gemma3:4b", "Gemma 系列主流视觉基线"],
            ["规则基线", "local-baseline", "无", "OpenCV 规则兜底"],
        ],
    )
    add_figure_placeholder(document, "图 4-3 Ollama 主模型与两个 4B 视觉基线验证截图。")

    add_heading_like(document, "五、自选无人机短视频素材")
    add_body(
        document,
        "本实验选择 Pexels 公开视频作为自选素材。原始视频文件为 upload_[REDACTED]_yubohan_drone.mp4，"
        "后续使用项目内 FFmpeg 压缩为 20 秒、720p 的 upload_[REDACTED]_yubohan_drone_review.mp4，"
        "以降低上传和模型抽帧耗时。"
    )
    add_body(
        document,
        "无人机视频适合作为审核场景，因为航拍画面具有高视角、广覆盖和连续运动特点，"
        "可能包含住宅、道路、车辆、人群、厂区、学校、基础设施或地理位置线索。平台发布前需要判断是否存在隐私、地理安全、飞行合规或画面质量问题。"
    )
    add_code(
        document,
        [
            "来源平台：Pexels public video",
            "直接媒体 URL：https://videos.pexels.com/video-files/7457249/7457249-uhd_3840_2160_25fps.mp4",
            "License：https://www.pexels.com/license/",
            "本地实验文件：data/incoming/upload_[REDACTED]_yubohan_drone_review.mp4",
        ],
    )
    add_figure_placeholder(document, "图 4-4 Pexels 无人机公开视频来源与许可说明截图。")

    add_heading_like(document, "六、端到端运行与严格验收")
    add_body(
        document,
        "严格验收时设置 ALLOW_LOCAL_MODEL_FALLBACK=0，表示如果本地多模态模型不可用，脚本必须失败，"
        "不能悄悄降级到 OpenCV baseline。verify_demo.py 已经在该设置下通过，输出 verification passed，"
        "并给出 published、review、rejected、processing 统计。"
    )
    add_code(
        document,
        [
            "$env:STUDENT_ID = \"[REDACTED]\"",
            "$env:ALLOW_LOCAL_MODEL_FALLBACK = \"0\"",
            ".\\.venv\\Scripts\\python.exe scripts\\create_demo_video.py",
            ".\\.venv\\Scripts\\python.exe scripts\\verify_demo.py",
        ],
    )
    add_body(
        document,
        "create_demo_video.py 在本作品集版本中只检查并列出已下载的 Pexels 真实公开视频，不再生成合成动画样本。"
    )
    add_figure_placeholder(document, "图 4-5 verify_demo.py 严格验收通过截图。")

    add_heading_like(document, "七、本机模型基线对照")
    add_body(
        document,
        "本实验对四组候选模型进行对照：Ministral 3 8B、Qwen3-VL 4B、Gemma 3 4B 和 OpenCV Local Baseline。"
        "为了避免只在正常素材上得到全 0 风险分，对照脚本同时覆盖正常无人机航拍、Pexels 真实夜间低照度街景和 Pexels 真实舞台灯光闪烁片段三个场景。"
        "这样既能比较模型语义标签和推理耗时，也能看到 published/review 与 0.0/42.0/48.0 风险分数的区分。"
    )
    add_table(
        document,
        [
            ["场景", "触发原因", "状态", "风险分数", "说明"],
            ["normal_drone", "policy_passed", "published", "0.0", "正常无人机居民区航拍，未命中高风险信号"],
            ["low_light_review", "too_dark", "review", "42.0", "画面过暗，自动理解置信度下降，需要人工复核"],
            ["flash_risk_review", "flash_risk, red_dominance, high_motion", "review", "48.0", "真实舞台灯光片段存在强亮度跳变和高运动幅度，可能造成观看不适"],
        ],
    )
    add_table(
        document,
        [
            ["模型", "后端", "正常无人机场景结果", "耗时", "主要语义标签"],
            ["Ministral 3 8B", "local_ollama_vlm", "published / 0.0", "约 9.41s", "居民区巡检、无人机拍摄、道路全景、住宅区监测"],
            ["Qwen3-VL 4B", "local_ollama_vlm", "published / 0.0", "约 13.40s", "英国郊区、住宅区、无人机航拍、城市景观"],
            ["Gemma 3 4B", "local_ollama_vlm", "published / 0.0", "约 4.85s", "住宅区、鸟瞰图、郊区、无人机"],
            ["OpenCV Local Baseline", "local_baseline", "published / 0.0", "约 2.11s", "横屏、自然光、运动明显等低层视觉标签"],
        ],
    )
    add_body(
        document,
        "可以看出，OpenCV baseline 运行最快，但只能给出亮度、运动、横竖屏等低层信号；"
        "三个本地视觉语言模型能够识别居民区、航拍、无人机巡检、住宅区等语义标签。"
        "主模型 ministral-3:8b 在本实验中输出更贴近无人机巡检业务描述，因此作为默认主链路模型。"
    )
    add_figure_placeholder(document, "图 4-6 本机模型基线对照输出截图。")

    add_heading_like(document, "八、Demo 网站上传与异步审核")
    add_body(
        document,
        "FastAPI 服务启动后访问 http://127.0.0.1:5050。上传标题统一为 [REDACTED]_Bohan Yu_自选上传_无人机巡检短视频。"
        "上传接口先返回 processing 状态，使视频卡片立即出现在页面中；随后 worker 在后台完成理解和审核。"
    )
    add_code(
        document,
        [
            ".\\.venv\\Scripts\\python.exe -m app.server",
            ".\\.venv\\Scripts\\python.exe scripts\\run_upload_api_evidence.py",
        ],
    )
    add_figure_placeholder(document, "图 4-7 FastAPI 服务启动截图。")
    add_figure_placeholder(document, "图 4-8 上传后页面立即显示 processing 视频卡片截图。")

    add_heading_like(document, "九、API 证据与事件追踪")
    add_body(
        document,
        "run_upload_api_evidence.py 自动保存 immediate 和 final 两类 API 证据。最终 /api/videos 中，自选无人机视频状态为 published，"
        "risk_score 为 0.0，模型 backend 为 local_ollama_vlm，selected_id 为 ministral-3-8b-ollama，ollama_model 为 ministral-3:8b。"
    )
    add_body(
        document,
        "事件日志记录了 ingest、stream、queued、worker、model_select、frame_sample、preprocess、vlm_understanding、understanding、moderation、publish 等阶段，"
        "证明系统具有完整的异步处理和可观测能力。"
    )
    add_figure_placeholder(document, "图 4-9 /api/videos 最终结果截图。")
    add_figure_placeholder(document, "图 4-10 /api/health 队列与模型状态截图。")
    add_figure_placeholder(document, "图 4-11 /api/events 全链路事件追踪截图。")

    add_heading_like(document, "十、最终审核结果分析")
    add_body(
        document,
        "本次无人机视频最终被判定为 published，风险分数为 0.0。模型摘要认为视频展示了居民区鸟瞰画面，"
        "包括道路、房屋、绿化带等基础设施，未见明显人群或车辆活动。审核规则未命中本实验中的高风险信号，因此允许自动发布。"
    )
    add_body(
        document,
        "但这并不代表无人机视频天然无需审核。真实平台中，无人机画面可能暴露个人隐私、车牌、人脸、住宅窗户、厂区设备、交通节点、学校、机场或能源设施。"
        "因此需要结合多模态理解、规则审核、地理围栏、OCR/目标检测和人工复核进行综合判断。"
    )
    add_figure_placeholder(document, "图 4-12 网页最终审核结果截图。")
    add_figure_placeholder(document, "图 4-13 模型选择区截图。")

    add_heading_like(document, "十一、Kafka 加分项说明")
    add_body(
        document,
        "本实验默认使用 SQLite jobs 表模拟消息队列，基础闭环不依赖 Kafka。为了与前三章流批一体实验衔接，"
        "系统提供 kafka_video_producer.py 和 kafka_ai_consumer.py，可将视频进入事件发送到 short_video_ingest_[REDACTED]，"
        "审核结果写入 short_video_result_[REDACTED]。本次已经确认 WSL Docker 中的 bigdata-kafka 容器可用，localhost:9092 端口可连通，"
        "并完成 producer、consumer 和 result topic 的端到端验证。"
    )
    add_body(
        document,
        "实际运行中，producer 向 short_video_ingest_[REDACTED] 发送标题为 [REDACTED]_Bohan Yu_Kafka_无人机巡检短视频 的无人机视频事件；"
        "consumer 处理后输出 processed 日志，并把 status=published、risk_score=0.0、tags 等结果写入 short_video_result_[REDACTED]。"
        "该扩展证明本地 SQLite 队列可以替换为真正的分布式消息队列。"
    )
    add_figure_placeholder(document, "图 4-14 Kafka topic 创建截图。")
    add_figure_placeholder(document, "图 4-15 Kafka producer/consumer/result 端到端结果截图。")

    add_heading_like(document, "十二、实验总结")
    add_body(
        document,
        "本实验完成了短视频上传、异步队列、多模态理解、风险审核、标签生成和发布展示的完整闭环。"
        "与同步等待模型推理相比，processing 状态和后台 worker 让视频能够先进入页面，提高了用户体验和系统稳定性。"
        "本地 8B VLM 能够识别无人机、居民区、鸟瞰、基础设施等语义信息，明显优于纯 OpenCV 规则基线。"
    )
    add_body(
        document,
        "如果部署到真实短视频平台，需要将本地媒体目录替换为 MinIO/S3，将 SQLite jobs 替换为 Kafka/Pulsar，"
        "将 SQLite 元数据替换为 PostgreSQL/ClickHouse/Paimon，将本地 Ollama 替换为可扩缩容的模型服务。"
        "同时还需要进一步加入失败重试、权限审计、模型监控、内容安全策略和人工复核流程。"
    )
    add_body(document, "本次实验追踪学号：[REDACTED]。")

    document.save(TARGET)
    print(TARGET)


if __name__ == "__main__":
    main()
