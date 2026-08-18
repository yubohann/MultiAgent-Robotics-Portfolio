"""Strengthen the Lab 4 report while preserving the base cover and captions."""

from pathlib import Path
import shutil

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
LAB_ROOT = PROJECT_ROOT.parents[1]
BASE = next(p for p in LAB_ROOT.glob("*.docx") if "before_strengthen" in p.name)
TARGET = next(
    p
    for p in LAB_ROOT.glob("*.docx")
    if "[REDACTED]" in p.name and "before" not in p.name and "backup" not in p.name
)
REPAIR_BACKUP = LAB_ROOT / "短视频审核实验报告_[REDACTED]_Bohan Yu_before_preserve_base_repair.docx"


def set_font(paragraph, size: int = 10) -> None:
    for run in paragraph.runs:
        run.font.name = "Microsoft YaHei"
        run.font.size = Pt(size)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def insert_body_after(document: Document, paragraph_index: int, text: str) -> None:
    paragraph = document.paragraphs[paragraph_index]
    new_paragraph = paragraph.insert_paragraph_before("")
    paragraph._p.addnext(new_paragraph._p)
    new_paragraph.paragraph_format.first_line_indent = Pt(20)
    new_paragraph.paragraph_format.line_spacing = 1.25
    run = new_paragraph.add_run(text)
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(10)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def find_paragraph(document: Document, startswith: str) -> int:
    for index, paragraph in enumerate(document.paragraphs):
        if paragraph.text.strip().startswith(startswith):
            return index
    raise ValueError(f"paragraph not found: {startswith}")


def strengthen(document: Document) -> None:
    """Insert additional analysis paragraphs without touching images/captions."""
    insert_body_after(
        document,
        find_paragraph(document, "本综合实验围绕实时短视频流"),
        "从业务角度看，短视频审核的难点不是单纯上传文件，而是在用户体验、审核准确性和系统吞吐之间取得平衡。"
        "如果上传接口同步等待多模态模型推理，用户会长时间卡在请求中；如果完全放弃模型理解，又无法解释画面主题、风险线索和标签来源。"
        "因此，本实验采用“上传立即可见、后台异步审核、事件全程可追踪”的方案，把用户交互链路与模型推理链路解耦。",
    )
    insert_body_after(
        document,
        find_paragraph(document, "本实验自选素材来自 Pexels"),
        "选择公开视频而不是随意上传网络视频，是为了同时满足版权合规、内容安全和结果可复现。"
        "无人机航拍用于主业务场景，夜间低照度和舞台灯光闪烁用于验证复核规则，运动场跑步用于证明正常真实视频可自动发布。"
        "这些素材分工使报告中的 published/review 结论具有场景区分度，而不是只在单一正常视频上得到全 0 风险分。",
    )
    insert_body_after(
        document,
        find_paragraph(document, "系统由前端页面、FastAPI API 服务"),
        "该架构对应真实短视频平台的常见分层：入口服务负责接收视频并生成初始状态，队列负责削峰和异步调度，媒体存储保存原始视频和关键帧，"
        "模型服务负责视觉语义理解，审核规则再把模型输出和可解释指标转成 published、review 或 rejected。"
        "本实验虽然是单机实现，但接口边界与工业系统一致，后续可以替换为 Kafka、MinIO、PostgreSQL/ClickHouse 和分布式模型服务。",
    )
    insert_body_after(
        document,
        find_paragraph(document, "对照脚本覆盖正常无人机航拍"),
        "模型对照的重点不只是速度比较，还包括语义能力比较。OpenCV baseline 能稳定计算亮度、运动和闪烁，适合作为兜底；"
        "但它无法真正理解“无人机巡检、居民区、夜间街景、舞台灯光”等语义。Qwen3-VL 4B 与 Gemma 3 4B 证明系统具备横向替换模型的能力，"
        "Ministral 3 8B 则作为主链路模型提供更完整的视觉理解结果。",
    )
    insert_body_after(
        document,
        find_paragraph(document, "随后启动 FastAPI 服务并访问"),
        "网页 processing 截图的意义是证明上传请求没有等待模型完成。视频先进入页面和本地队列，摘要、标签和风险分数由 worker 异步补齐。"
        "这种设计更接近真实平台：用户侧关注上传是否成功，平台侧通过后台审核逐步补充理解结果和发布状态。",
    )
    insert_body_after(
        document,
        find_paragraph(document, "最终无人机视频被判定为 published"),
        "无人机视频最终自动发布，并不代表无人机素材不需要审核。相反，无人机画面具有高视角和广覆盖特征，容易同时拍到道路、住宅、车辆、基础设施和园区边界。"
        "本次素材未命中高风险规则，因此 risk_score 为 0.0；但系统仍保留关键帧理解、审核原因和事件日志，保证后续可解释和可追溯。",
    )
    insert_body_after(
        document,
        find_paragraph(document, "Kafka producer 发送一条无人机短视频进入事件"),
        "Kafka 加分项只发送一条无人机视频事件，是为了突出“视频进入事件 -> AI 审核消费者 -> result topic 输出”的链路验证。"
        "低照度、舞台闪烁和运动场视频已经在 FastAPI/SQLite 主链路和 baseline 对照中验证；Kafka 部分用于证明该审核逻辑可以从单机队列迁移到课程前三个实验使用的流处理思想。",
    )
    insert_body_after(
        document,
        find_paragraph(document, "本实验的关键价值在于把短视频审核拆成"),
        "与前三个实验相比，本实验把 Kafka/Flink/Paimon 中的流式处理、状态观测和数据治理思想扩展到了非结构化媒体场景。"
        "视频不再只是普通文件，而是带有事件、队列状态、模型理解结果和审核决策的流式对象。"
        "这体现了云计算与大数据处理课程从结构化数据流到多模态智能应用的延伸。",
    )
    insert_body_after(
        document,
        find_paragraph(document, "本次综合实验完成了实时短视频上传"),
        "本实验仍存在局限：关键帧抽样可能漏掉极短暂的人脸、车牌或敏感标识；VLM 对远距离小目标和地理位置判断能力有限；教学版规则分数也不能直接等同生产策略。"
        "生产环境还应接入 OCR、人脸/车牌检测、地理围栏、人工复核队列、模型评测集和灰度发布机制，以降低误判和漏判风险。",
    )


def main() -> None:
    if not REPAIR_BACKUP.exists() and TARGET.exists():
        shutil.copy2(TARGET, REPAIR_BACKUP)
    document = Document(str(BASE))
    strengthen(document)
    document.save(str(TARGET))
    print(TARGET)
    print(f"base: {BASE}")
    print(f"backup: {REPAIR_BACKUP}")


if __name__ == "__main__":
    main()
